"""
Export AuthoritySection for a case to a Word (.docx) file.
"""

from __future__ import annotations

import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from scripts.link_case_to_law import link_ds_3960

from knowledge.legal_query import LegalQuery
from knowledge.models.authority_section import build_authority_section


def export_authorities_docx(
    path: Path,
    *,
    case_id: str = "case:DS.3960.2025",
) -> Path:
    try:
        from docx import Document
        from docx.enum.text import WD_LINE_SPACING
        from docx.shared import Pt
    except ImportError as exc:
        raise ImportError("python-docx is required: pip install python-docx") from exc

    graph, linked_id = link_ds_3960()
    cid = linked_id or case_id
    section = build_authority_section(LegalQuery(graph), cid)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    def add_line(text: str = "", *, bold: bool = False) -> None:
        para = doc.add_paragraph()
        run = para.add_run(text)
        run.bold = bold
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        pf = para.paragraph_format
        pf.space_after = Pt(0)
        pf.space_before = Pt(0)
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE

    add_line("PODSTAWA PRAWNA I ORZECZNICTWO", bold=True)
    add_line()
    add_line("Przepisy, na których oparto stanowisko:", bold=True)
    if not section.statutes:
        add_line("1. (brak)")
    else:
        for i, item in enumerate(section.statutes, 1):
            add_line(f"{i}. {item.title}")
            if item.summary:
                add_line(f"   {item.summary}")
    add_line()
    add_line("Orzecznictwo wspierające ocenę prawną:", bold=True)
    if not section.case_law:
        add_line("1. (brak)")
    else:
        for i, item in enumerate(section.case_law, 1):
            add_line(f"{i}. {item.citation or item.title}")
            if item.summary:
                add_line(f"   Teza: {item.summary}")
    add_line()
    add_line("Wybrane interpretacje (przepis wiodący):", bold=True)
    if not section.interpretations:
        add_line("— (brak)")
    else:
        for item in section.interpretations:
            if item.summary:
                add_line(f"— {item.title}: {item.summary}")
            else:
                add_line(f"— {item.title}")

    path = Path(path)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))
    return path.resolve()


def main() -> None:
    out = Path("output") / "cases" / "DS_3960_2025" / "authorities_from_graph.docx"
    saved = export_authorities_docx(out)
    print("Saved:", saved)


if __name__ == "__main__":
    main()